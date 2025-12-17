"""
Multi-Format Support - Support for Word (DOCX) and Excel (XLSX) documents.
"""
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd
import structlog

logger = structlog.get_logger()

# Check for optional dependencies
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word support disabled.")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available. Excel support may be limited.")


class WordParser:
    """Parse Microsoft Word documents (.docx)."""
    
    def __init__(self):
        """Initialize Word parser."""
        self.logger = logger.bind(component="WordParser")
        self.available = DOCX_AVAILABLE
    
    def is_available(self) -> bool:
        """Check if Word parsing is available."""
        return self.available
    
    def extract_text(self, docx_path: str) -> str:
        """Extract text from Word document.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        if not self.available:
            raise RuntimeError(
                "Word support not available. Install: pip install python-docx"
            )
        
        self.logger.info("Extracting text from Word document", file=docx_path)
        
        try:
            doc = DocxDocument(docx_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs]
            text = '\n'.join(paragraphs)
            
            self.logger.info(
                "Text extraction completed",
                paragraphs=len(paragraphs),
                text_length=len(text)
            )
            
            return text
            
        except Exception as e:
            self.logger.error("Word extraction failed", error=str(e))
            raise
    
    def extract_tables(self, docx_path: str) -> List[pd.DataFrame]:
        """Extract tables from Word document.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            List of DataFrames
        """
        if not self.available:
            raise RuntimeError("Word support not available")
        
        try:
            doc = DocxDocument(docx_path)
            
            tables = []
            
            for table in doc.tables:
                # Extract table data
                data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    data.append(row_data)
                
                if data:
                    # Use first row as header
                    df = pd.DataFrame(data[1:], columns=data[0])
                    tables.append(df)
            
            self.logger.info("Tables extracted", count=len(tables))
            
            return tables
            
        except Exception as e:
            self.logger.error("Table extraction failed", error=str(e))
            raise
    
    def extract_metadata(self, docx_path: str) -> Dict[str, Any]:
        """Extract document metadata.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary of metadata
        """
        if not self.available:
            return {}
        
        try:
            doc = DocxDocument(docx_path)
            core_props = doc.core_properties
            
            return {
                'author': core_props.author,
                'title': core_props.title,
                'subject': core_props.subject,
                'created': core_props.created,
                'modified': core_props.modified,
                'revision': core_props.revision
            }
        except Exception as e:
            self.logger.error("Metadata extraction failed", error=str(e))
            return {}


class ExcelParser:
    """Parse Microsoft Excel documents (.xlsx, .xls)."""
    
    def __init__(self):
        """Initialize Excel parser."""
        self.logger = logger.bind(component="ExcelParser")
    
    def extract_sheets(self, excel_path: str) -> Dict[str, pd.DataFrame]:
        """Extract all sheets from Excel file.
        
        Args:
            excel_path: Path to Excel file
            
        Returns:
            Dictionary mapping sheet names to DataFrames
        """
        self.logger.info("Extracting sheets from Excel", file=excel_path)
        
        try:
            # Read all sheets
            sheets = pd.read_excel(excel_path, sheet_name=None, engine='openpyxl')
            
            self.logger.info("Sheets extracted", count=len(sheets))
            
            return sheets
            
        except Exception as e:
            self.logger.error("Excel extraction failed", error=str(e))
            raise
    
    def extract_sheet(
        self,
        excel_path: str,
        sheet_name: Optional[str] = None
    ) -> pd.DataFrame:
        """Extract specific sheet from Excel file.
        
        Args:
            excel_path: Path to Excel file
            sheet_name: Sheet name (None for first sheet)
            
        Returns:
            DataFrame
        """
        try:
            df = pd.read_excel(
                excel_path,
                sheet_name=sheet_name or 0,
                engine='openpyxl'
            )
            
            return df
            
        except Exception as e:
            self.logger.error("Sheet extraction failed", error=str(e))
            raise
    
    def get_sheet_names(self, excel_path: str) -> List[str]:
        """Get list of sheet names.
        
        Args:
            excel_path: Path to Excel file
            
        Returns:
            List of sheet names
        """
        try:
            excel_file = pd.ExcelFile(excel_path, engine='openpyxl')
            return excel_file.sheet_names
        except Exception as e:
            self.logger.error("Failed to get sheet names", error=str(e))
            return []


class CSVParser:
    """Parse CSV files."""
    
    def __init__(self):
        """Initialize CSV parser."""
        self.logger = logger.bind(component="CSVParser")
    
    def extract_data(
        self,
        csv_path: str,
        encoding: str = 'utf-8',
        delimiter: str = ','
    ) -> pd.DataFrame:
        """Extract data from CSV file.
        
        Args:
            csv_path: Path to CSV file
            encoding: File encoding
            delimiter: CSV delimiter
            
        Returns:
            DataFrame
        """
        self.logger.info("Extracting CSV data", file=csv_path)
        
        try:
            df = pd.read_csv(
                csv_path,
                encoding=encoding,
                delimiter=delimiter
            )
            
            self.logger.info("CSV data extracted", rows=len(df), columns=len(df.columns))
            
            return df
            
        except Exception as e:
            self.logger.error("CSV extraction failed", error=str(e))
            raise
    
    def detect_delimiter(self, csv_path: str) -> str:
        """Detect CSV delimiter.
        
        Args:
            csv_path: Path to CSV file
            
        Returns:
            Detected delimiter
        """
        import csv
        
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                return delimiter
        except Exception:
            return ','  # Default to comma
